#!/usr/bin/python
# -*- coding: cp1251 -*-
import docx
import docx2txt
import pypandoc
import os
import time
import numpy as np
import main_settings
from threading import Thread
from PIL import Image
from pathlib import Path
from docx.oxml import OxmlElement, ns
from docx2python import docx2python
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT

# После логирования программа не выносит нижний колонтитул

start_time = time.time()

class LoggingCallback:

    def __init__(self):
        self.text = []
        self.header_text = [] 
        self.num_footnotes = []
    
    def __call__(self, text, header_text, num_footnotes):
        self.text.append(text)
        self.header_text.append(header_text)
        self.num_footnotes.append(num_footnotes)

class NewDocument:

    def __init__(self, filename):
        self.filename = filename

    @staticmethod
    def find_words(words):
        if len(words) > 0:
            for word in words:
                if word.lower() == 'из':
                    pass
    
    @staticmethod
    def flattenlist(nestedlist): 
        if len(nestedlist) == 0: 
            return nestedlist 
        if isinstance(nestedlist[0], list): 
            return NewDocument.flattenlist(nestedlist[0]) + NewDocument.flattenlist(nestedlist[1:]) 
        return nestedlist[:1] + NewDocument.flattenlist(nestedlist[1:])        

    def adding_margins(self):
        table_word, hyperlinks_text, foot_flag = [], [], False
        document = docx2python(Path("Documents") / str(self.filename))
        for obj in document.body:
            if len(obj) > 1 and len(obj[0][0]) == 1 and obj[0][0] != ['\t']:
                numbers = [
                    [obj[_][number][0] for number in range(
                    len(obj[0]))] for _ in range(len(obj))
                ]
                for main in numbers:
                    for letter in main:
                        table_word.append(letter)
                callback(numbers, numbers, 0)
            else:
                new_text = self.flattenlist(obj)
                for line in range(len(new_text)):
                    if new_text[line] != '':
                        if new_text[line].find('</a>') != -1:
                            hyperlinks_text.append(
                                new_text[line][: new_text[line].index('<a')] +
                                new_text[line][new_text[line].index('">') + 2: new_text[line].index('</a>')] +
                                new_text[line][new_text[line].index('</a>') + 4:]
                            )
                        if new_text[line].find('.png') != -1:
                            callback(
                                new_text[line][new_text[line].find(
                                'media') + 6:][: new_text[line][new_text[line].find('media') + 6:].find('----')],
                                new_text[line][new_text[line].find(
                                'media') + 6:][: new_text[line][new_text[line].find('media') + 6:].find('----')], 0
                            )
                        elif new_text[line].find('footnote') != -1:
                            if new_text[line].find('\t') != -1:
                                phrase = list(
                                    new_text[line][new_text[line].find(
                                    '\t') + 1: new_text[line].find('----')]
                                )
                            else:
                                phrase = list(
                                    new_text[line][:new_text[line].find('----')])
                            while True:
                                if phrase[0].isalpha():
                                    break
                                if phrase[0] == ' ':
                                    del phrase[0]
                            if new_text[line].find('\t') != -1:
                                phrase = np.insert(phrase, 0, '\t')
                                phrase = np.insert(phrase, 0, '-')
                            else:
                                for _ in range(12):
                                    phrase = np.insert(phrase, 0, ' ')
                            phrase = ''.join(phrase)
                            callback(phrase, phrase, 1)
                            foot_flag = True
                        else:
                            phrase = list(new_text[line])
                            if new_text[line].find('--\t') != -1 and new_text[line] != '\t':
                                callback(new_text[line][new_text[line].find('--\t') + 1:], 
                                    new_text[line][new_text[line].find('--\t') + 3:],
                                    0)
                                continue
                            if len(new_text[line].split()) <= 10:
                                callback(new_text[line].capitalize(), new_text[line].capitalize(), 0)
                                continue
                            while True:
                                if phrase[0].isalpha():
                                    break
                                if phrase[0] == ' ':
                                    del phrase[0]
                            table_word.append(''.join(phrase))
                            for _ in range(12):
                                phrase = np.insert(phrase, 0, ' ')
                            phrase = ''.join(phrase)
                            callback(phrase, new_text[line], 0)
        # for my in callback.text:
        #     print(my)
        #     table_word.append(' '.join(my.split()))
        return foot_flag, table_word, hyperlinks_text

    def adding_footnotes(self):
        with open(os.path.join("Documents", str(self.filename[:self.filename.find('.docx')]) + '.txt'), 'r', encoding='utf8') as file:
            f = file.readlines()
            note = []
            for word in range(len(f)):
                sp, summa_end = word, 0
                if f[word].find('footnote') != -1:
                    summa_end, flag_https = 0, False
                    while summa_end != 2 or word < len(f):
                        if f[word].find('https') != -1:
                            flag_https = True
                        if f[word].find(']]') != -1:
                            notes += f[word][: f[word].index(']') + 1]
                            break
                        elif f[word].find(']') != -1:
                            summa_end += 1
                            if summa_end == 1 and not flag_https:
                                notes += f[word][: f[word].index(']')]
                                break
                            if summa_end != 2 and flag_https:
                                notes += f[word][: f[word].index('\n')] + ' '
                            else:
                                notes += f[word][: f[word].index(']')]
                                break
                        else:
                            if f[word].find('[') != -1:
                                notes = f[word][f[word].index(
                                    '[') + 1: f[word].index('\n')] + ' '
                            else:
                                notes += f[word][: f[word].index('\n')] + ' '
                        word += 1
                        if word == len(f):
                            break
                    word = sp
                    note.append(notes)
            footnotes = dict((number, par) for number, par in enumerate(note))
        os.remove(os.path.join("Documents", str(
            self.filename[:self.filename.find('.docx')]) + '.txt'))
        return footnotes

    def adding_heading(self, name):
        if callback.text[name] == '' and headings[name -
                                            1] == 0 and name + 1 <= len(headings):
            try:
                headings[name + 1] == 0
            except IndexError:
                return
            else:
                return
        if headings[name] and callback.text[name - 1] != '' and name != 0:
            paragrahp = new_doc.add_paragraph('')
        paragrahp = new_doc.add_paragraph('')
        self.find_words(callback.text[name].split())
        run = paragrahp.add_run(callback.text[name])
        run.bold = True
        run.font.size = Pt(main_settings.pt + 2)
        p_fmt = paragrahp.paragraph_format
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fmt.line_spacing = 1.5
        p_fmt.space_before = Pt(1.5)
        p_fmt.space_after = Pt(1.5)

    def adding_paragraph(self, name):
        if name != 0:
            if tables[name + 1] == 1:
                new_doc.add_paragraph(' ')
                self.find_words(callback.text[name].split())
                paragrahp = new_doc.add_paragraph('')
                run = paragrahp.add_run(callback.text[name])
                run.italic = True
                run.font.size = Pt(main_settings.pt - 2)
                p_fmt = paragrahp.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif pictures[name - 1] == 1 and callback.text[name] != ' ':
                self.find_words(callback.text[name].split())
                paragrahp = new_doc.add_paragraph('')
                run = paragrahp.add_run(callback.text[name])
                run.italic = True
                run.font.size = Pt(main_settings.pt - 2)
                p_fmt = paragrahp.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                self.find_words(callback.text[name].split())
                paragrahp = new_doc.add_paragraph(callback.text[name])
                p_fmt = paragrahp.paragraph_format
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_fmt.line_spacing = main_settings.line_spacing
            p_fmt.space_before = Pt(main_settings.line_spacing)
            p_fmt.space_after = Pt(main_settings.line_spacing)
        return paragrahp

    def adding_picture(self, name):
        with Image.open(Path("Documents") / str(callback.text[name])) as img:
            width, height = img.size
            img.save(Path("Documents") /
                     str(callback.text[name]), dpi=(800, 800), optimize=True, quality=100)
        if width > 700:
            new_doc.add_picture(os.path.join("Documents", str(
                callback.text[name])), width=Mm(width / 4.25), height=Mm(height / 4.25))
        else:
            new_doc.add_picture(os.path.join("Documents", str(
                callback.text[name])), width=Mm(width / 3.5), height=Mm(height / 3.5))
        new_doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        os.remove(Path("Documents") / str(callback.text[name]))

    def adding_list(self, name):
        if callback.text[name].find('-\t') != -1:
            self.find_words(
                callback.text[name][callback.text[name].find('\t') + 1:].split())
            paragrahp = new_doc.add_paragraph(
                callback.text[name][callback.text[name].find('\t') + 1:], style='List Bullet')
        else:
            self.find_words(
                callback.text[name][callback.text[name].find('\t') + 1:].split())
            paragrahp = new_doc.add_paragraph(
                callback.text[name][callback.text[name].find('\t') + 1:], style='List Number')
        p_fmt = paragrahp.paragraph_format
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_fmt.line_spacing = main_settings.line_spacing
        p_fmt.space_before = Pt(main_settings.line_spacing)
        p_fmt.space_after = Pt(main_settings.line_spacing)
        return paragrahp

    def adding_table(self, txt, name):
        doc = docx.Document(Path("Documents") / str(self.filename))
        unique, merged, max_row, max_col, added, cells_text, all_cells_text = [], [], np.array([]), np.array([]), [], [], []
        for row in doc.tables[table_num].rows:
            info = []
            for cell in row.cells:
                tc = cell._tc
                max_row, max_col = np.append(
                    tc.bottom, max_row), np.append(tc.right, max_col)
                cell_loc = (tc.top, tc.bottom, tc.left, tc.right)
                if tc.bottom - tc.top > 1:
                    if cell_loc not in merged:
                        info.append(cell.text)
                        all_cells_text.extend(cell.text.split('\n'))
                    else:
                        info.append(' ')
                        all_cells_text.append(' ')
                elif tc.right - tc.left > 1:
                    if cell_loc not in merged:
                        info.append(cell.text)
                        all_cells_text.extend(cell.text.split('\n'))
                    else:
                        info.append(' ')
                        all_cells_text.append(' ')
                else:
                    info.append(cell.text)
                    all_cells_text.extend(cell.text.split('\n'))
                if tc.bottom - tc.top > 1 or tc.right - tc.left > 1 and cell_loc not in merged:
                    merged.append(cell_loc)
                else:
                    unique.append(cell_loc)
            cells_text.append(info)
        table = new_doc.add_table(rows=0, cols=int(
            np.amax(max_col)), style='Table Grid')
        for row in range(int(np.amax(max_row))):
            table.add_row().cells
            for col in range(int(np.amax(max_col))):
                cell_paragraphs = [
                    paragraph for paragraph in table.cell(row, col).paragraphs]
                for paragraph in cell_paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                    paragraph._p = paragraph._element = None
                table.cell(
                    row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if row == 0:
                    p = table.cell(row, col).add_paragraph('')
                    if cells_text[row][col] != ' ':
                        self.find_words(cells_text[row][col])
                        run = p.add_run(cells_text[row][col])
                        if len(txt[name][0]) == 1:
                            run.font.size = Pt(main_settings.pt)
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            run.bold = True
                            run.font.size = Pt(main_settings.pt + 2)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    if cells_text[row][col] != '':
                        self.find_words(cells_text[row][col])
                        p = table.cell(row, col).add_paragraph(
                            cells_text[row][col])
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if tables[name + 1] == 1:
            new_doc.add_paragraph(' ')
        for row in new_doc.tables[table_num].rows:
            for cell in row.cells:
                tc = cell._tc
                loc = (tc.top, tc.bottom, tc.left, tc.right)
                if loc not in unique:
                    added.append(loc)
        for merge in merged:
            for cell_1 in range(len(added) - 1):
                for cell_2 in range(cell_1 + 1, len(added)):
                    if (added[cell_1][0] == merge[0] and added[cell_1][1] == merge[1] and added[cell_1][2] == merge[2] and added[cell_2][3] == merge[3] and added[cell_2][0] == merge[0] and added[cell_2][1] == merge[1]) or (
                            added[cell_2][0] == merge[0] and added[cell_2][1] == merge[1] and added[cell_2][2] == merge[2] and added[cell_1][3] == merge[3] and added[cell_1][0] == merge[0] and added[cell_1][1] == merge[1]) or (
                            added[cell_1][1] == merge[1] and added[cell_1][2] == merge[2] and added[cell_1][3] == merge[3] and added[cell_2][0] == merge[0] and added[cell_2][2] == merge[2] and added[cell_2][3] == merge[3]) or (
                            added[cell_2][1] == merge[1] and added[cell_2][2] == merge[2] and added[cell_2][3] == merge[3] and added[cell_1][0] == merge[0] and added[cell_1][2] == merge[2] and added[cell_1][3] == merge[3]):
                        table.cell(min([added[cell_1][0], added[cell_1][1]]), min([added[cell_1][2], added[cell_1][3]])).merge(
                            table.cell(min([added[cell_2][0], added[cell_2][1]]), min([added[cell_2][2], added[cell_2][3]])))
        for row in range(int(np.amax(max_row))):
            for col in range(int(np.amax(max_col))):
                cell_paragraphs = [
                    paragraph for paragraph in table.cell(row, col).paragraphs]
                for paragraph in cell_paragraphs:
                    if ' ' in paragraph.text:
                        p = paragraph._element
                        p.getparent().remove(p)
                        paragraph._p = paragraph._element = None
        return all_cells_text

    def adding_headers_and_footers(self, up, down):
        if len(up) > 0:
            header = new_doc.sections[0].header
            header_para = header.add_paragraph('')
            run = header_para.add_run(str(up[0]))
            run.italic = True
            run.font.size = Pt(main_settings.pt)
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if len(down) > 0:
            footer = new_doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = str(down[0]) + '\n\n'
            footer_para.runs[0].italic = True
            footer_para.runs[0].font.size = Pt(main_settings.pt)

    def add_page_number(self, paragraph):

        def create_attribute(element, name, value):
            element.set(ns.qn(name), value)

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        page_num_run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
        instrText = OxmlElement('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')
        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

    def adding_hyperlink(self, paragraph, color, underline):
        url = paragraph.text[paragraph.text.index(
            'https'): paragraph.text.index('>') - 1]
        t = paragraph.text[paragraph.text.index(
            '">') + 2: paragraph.text.index('</a>')]
        if list_flag_num or list_flag_bul:
            ph_1 = paragraph.text[paragraph.text.find(
                '\t') + 1: paragraph.text.index('<a')]
        else:
            ph_1 = paragraph.text[: paragraph.text.index('<a')]
        ph_2 = paragraph.text[paragraph.text.index('</a>') + 4:]
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
        if list_flag_num:
            paragraph = new_doc.add_paragraph('', style='List Number')
        elif list_flag_bul:
            paragraph = new_doc.add_paragraph('', style='List Bullet')
        else:
            paragraph = new_doc.add_paragraph('')
        p_fmt = paragraph.paragraph_format
        if italic_flag:
            run = paragraph.add_run(ph_1)
            run.italic = True
            run.font.size = Pt(main_settings.pt - 2)
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.add_run(ph_1)
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_fmt.line_spacing = main_settings.line_spacing
        p_fmt.space_before = Pt(main_settings.line_spacing)
        p_fmt.space_after = Pt(main_settings.line_spacing)
        part = paragraph.part
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        if color is not None:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)
        new_run.append(rPr)
        new_run.text = t
        hyperlink.append(new_run)
        run = paragraph.add_run()
        run._r.append(hyperlink)
        if italic_flag:
            run.italic = True
            run.font.size = Pt(main_settings.pt - 2)
        run.font.underline = underline
        if italic_flag:
            run = paragraph.add_run(ph_2)
            run.italic = True
            run.font.size = Pt(main_settings.pt - 2)
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.add_run(ph_2)
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_fmt.line_spacing = main_settings.line_spacing
        p_fmt.space_before = Pt(main_settings.line_spacing)
        p_fmt.space_after = Pt(main_settings.line_spacing)
        a = ph_1 + t + ph_2
        return hyperlink, a

    def adding_hyperlink_in_footers(self, paragraph, color, underline):
        url = footnotes[summa][footnotes[summa].find(
            'https'): footnotes[summa].find('[')]
        t = footnotes[summa][footnotes[summa].find(
            '[') + 1: footnotes[summa].find(']')]
        part = paragraph.part
        ph_1 = footnotes[summa][: footnotes[summa].find('https')]
        ph_2 = footnotes[summa][footnotes[summa].find(']') + 1:]
        tables_text.append(ph_1 + t + ph_2)
        run = paragraph.add_run('[' + ph_1)
        run.bold = True
        run.font.size = Pt(10)
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)
        new_run.append(rPr)
        new_run.text = t
        hyperlink.append(new_run)
        run = paragraph.add_run()
        run._r.append(hyperlink)
        run.font.underline = underline
        run.bold = True
        run.font.size = Pt(10)
        run = paragraph.add_run(ph_2 + ']')
        run.bold = True
        run.font.size = Pt(10)


def delete_files():
    directory = Path("New Documents")
    if os.path.isdir("New Documents"):
        for filename in os.listdir(directory):
            f = os.path.join(directory, filename)
            if os.path.isfile(f) and filename.endswith('.docx'):
                os.remove(f)


delete_files()
directory, files = os.path.abspath("Documents"), []
assert len(os.listdir(directory)) > 0, 'Download files in the folder'
if not os.path.isdir("New Documents"):
    os.mkdir("New Documents")
for filename in os.listdir(directory):
    f = os.path.join(directory, filename)
    if os.path.isfile(f) and filename.endswith('.docx'):
        callback = LoggingCallback()
        files.append(filename[: filename.find('.docx')])
        new_documents = NewDocument(filename)
        text = new_documents.adding_margins()
        if text[0]:
            docxFilename = os.path.join("Documents", str(filename))
            pypandoc.convert_file(docxFilename, to='asciidoc', outputfile=os.path.join(
                "Documents", str(filename[:filename.find('.docx')]) + '.txt'))
            # Codecs: asciidoc, asciidoctor, beamer, biblatex, bibtex, commonmark, commonmark_x, context, csljson, docbook,
            # docbook4, docbook5, docx, dokuwiki, dzslides, epub, epub2, epub3, fb2, gfm, haddock, html, html4, html5, icml, ipynb,
            # jats, jats_archiving, jats_articleauthoring, jats_publishing, jira, json, latex, man, markdown, markdown_github,
            # markdown_mmd, markdown_phpextra, markdown_strict, markua, mediawiki, ms, muse, native, odt, opendocument, opml,
            # org, pdf, plain, pptx, revealjs, rst, rtf, s5, slideous, slidy, tei, texinfo, textile, xwiki, zimwiki
            footnotes = new_documents.adding_footnotes()
        my_doc = docx2txt.process(
            Path("Documents") / str(filename), directory).split('\n')
        new_doc = docx.Document()
        new_doc.sections[0].orientation = WD_ORIENT.PORTRAIT
        new_doc.sections[0].left_margin = Mm(main_settings.left_margin)
        new_doc.sections[0].right_margin = Mm(main_settings.right_margin)
        new_doc.sections[0].top_margin = Mm(main_settings.top_margin)
        new_doc.sections[0].bottom_margin = Mm(main_settings.bottom_margin)
        new_doc.styles['Normal'].font.name = main_settings.font_name
        new_doc.styles['Normal'].font.size = Pt(main_settings.pt)
        headings = [
            1 if not isinstance(
                callback.text[j],
                list) and len(
                callback.text[j].split()) <= 10 and callback.text[j] != '' and callback.text[j].find('\t') == -
            1 else 0 for j in range(len(callback.text))]
        pictures, tables_text, summa, table_num = [1 if not isinstance(callback.text[phrase], list) and callback.text[phrase].find(
            '.png') != -1 else 0 for phrase in range(len(callback.text))], [], 0, -1
        tables, hyper_text = [1 if isinstance(
            callback.text[j], list) else 0 for j in range(len(callback.text))], []
        tables.append(0)
        for phrase in range(len(callback.text)):
            list_flag_num, list_flag_bul, italic_flag, bold_flag = False, False, False, False
            if isinstance(callback.text[phrase], list):
                table_num += 1
                tables_text.extend(new_documents.adding_table(callback.text, phrase))
            elif callback.text[phrase].find('.png') != -1:
                new_documents.adding_picture(phrase)
            elif callback.text[phrase].find('\t') != -1:
                tables_text.append(' '.join(callback.text[phrase][2:].split()))
                if callback.text[phrase].find('</a>') != -1:
                    if callback.text[phrase].find('-\t') != -1:
                        list_flag_bul = True
                    else:
                        list_flag_num = True
                    paragrahp = new_doc.add_paragraph(callback.text[phrase])
                    hyperlink = new_documents.adding_hyperlink(
                        paragrahp, main_settings.hyperlink_color, main_settings.hyperlink_underline)
                    hyper_text.append(' '.join(hyperlink[-1].split()))
                else:
                    paragrahp = new_documents.adding_list(phrase)
                if callback.num_footnotes[phrase]:
                    if footnotes[summa].find('https') != -1:
                        new_documents.adding_hyperlink_in_footers(
                            paragrahp, main_settings.hyperlink_color, main_settings.hyperlink_underline)
                    else:
                        run = paragrahp.add_run(
                            ' [' + str(footnotes[summa]) + ']')
                        run.bold = True
                        run.font.size = Pt(10)
                    summa += 1
            elif len(callback.text[phrase].split()) <= 10:
                new_documents.adding_heading(phrase)
            else:
                if tables[phrase +
                          1] == 1 or pictures[phrase -
                                              1] == 1 and callback.text[phrase] != ' ':
                    italic_flag = True
                if callback.text[phrase].find('</a>') != -1:
                    paragraph = new_documents.adding_paragraph(phrase)
                    hyperlink = new_documents.adding_hyperlink(
                        paragraph, main_settings.hyperlink_color, main_settings.hyperlink_underline)
                    hyper_text.append(' '.join(hyperlink[-1].split()))
                else:
                    paragraph = new_documents.adding_paragraph(phrase)
                if callback.num_footnotes[phrase]:
                    if footnotes[summa].find('https') != -1:
                        new_documents.adding_hyperlink_in_footers(
                            paragraph, main_settings.hyperlink_color, main_settings.hyperlink_underline)
                    else:
                        run = paragraph.add_run(
                            ' [' + str(footnotes[summa]) + ']')
                        run.bold = True
                        run.font.size = Pt(10)
                    summa += 1
        upper_header_list, lower_header_list, header_flag = [], [], False
        for split in my_doc:
            if split != '':
                if split == callback.text:
                    break
                else:
                    upper_header_list.append(split)
        for split in my_doc:
            split = ' '.join(split.split())
            if split != '':
                if split.endswith(' '):
                    split = split[: - 2]
                if split not in callback.header_text and split not in text[-2] and split not in text[-1] and split not in upper_header_list and split not in tables_text and split not in callback.text and split not in hyper_text:
                    lower_header_list.append(split)
        if len(upper_header_list) > 0 or len(lower_header_list) > 0:
            header_flag = True
        if header_flag:
            new_documents.adding_headers_and_footers(
                upper_header_list, lower_header_list)
        new_documents.add_page_number(new_doc.sections[0].footer.paragraphs[0])
        save_name = Path("New Documents") / str(filename)
        new_doc.save(save_name)

end_time = time.time()
print(end_time - start_time)
