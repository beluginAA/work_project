import os
import time
import docx
import docx2txt
import shutil
import main_settings

from pathlib import Path
import numpy as np
from docx2python import docx2python
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from contextlib import suppress
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Mm
from PIL import Image

start_time = time.time()

class NewDocuments:

    def __init__(self, filename):
        self.filename = filename
        self.text = []
        self.header_text = []
        self.num_footnotes = []
    
    def __call__(self, text):
        self.text.append(text)  

    @staticmethod
    def flattenlist(nestedlist): 
        if len(nestedlist) == 0: 
            return nestedlist 
        if isinstance(nestedlist[0], list): 
            return NewDocuments.flattenlist(nestedlist[0]) + NewDocuments.flattenlist(nestedlist[1:]) 
        return nestedlist[:1] + NewDocuments.flattenlist(nestedlist[1:]) 

    def adding_margins(self): 
        
        text, links, headers, images = [], [], [], []
        document = docx2python(directory_file)
        for obj in document.body:
            if len(obj) > 1:
                new_document('TABLE')
            else:
                for phrase in self.flattenlist(obj):
                    if not phrase:
                        new_document(phrase)
                    elif phrase.find('.png') != -1:
                        new_document('IMAGE')
                        images.append(phrase)
                    elif phrase.find('</a>') != -1:
                        new_document('LINK')
                        links.append(phrase)
                    elif phrase[0] == '\t' or len(phrase.split()) < 10:
                        new_document('HEADER')
                        headers.append(phrase)
                    else:
                        phrase = ' ' * 14 + phrase.lstrip()
                        new_document('TEXT')
                        main_text.append(phrase)
        return text, links, headers, images
    
    def adding_picture(self, picture):
        with Image.open(os.path.join(directory, picture)) as img:
            width, height = img.size
            img.save(os.path.join(directory, picture), dpi=(800, 800), optimize=True, quality=100)
        if width > 700:
            width_mm, height_mm = width / 4.25, height / 4.25
        else:
            width_mm, height_mm = width / 3.5, height / 3.5
        new_doc.add_picture(os.path.join(directory, picture), width=Mm(width_mm), height=Mm(height_mm)).alignment = WD_ALIGN_PARAGRAPH.CENTER
        with suppress(FileNotFoundError):
            os.remove(os.path.join(directory, picture))

    def adding_heading(self, header):
        paragrahp = new_doc.add_paragraph('')
        # self.find_words(callback.text[name].split())
        run = paragrahp.add_run(header)
        run.bold = True
        run.font.size = Pt(main_settings.pt)
        p_fmt = paragrahp.paragraph_format
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fmt.line_spacing = 1
        p_fmt.space_before = Pt(1)
        p_fmt.space_after = Pt(1)

    def adding_paragraph(self, paragraph):
        if not paragrahp:
            # self.find_words(callback.text[name].split())
            paragrahp = new_doc.add_paragraph(paragrahp)
            p_fmt = paragrahp.paragraph_format
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_fmt.line_spacing = main_settings.line_spacing
            p_fmt.space_before = Pt(main_settings.line_spacing)
            p_fmt.space_after = Pt(main_settings.line_spacing)
        return paragrahp
    
    def add_page_number(self, paragraph):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        page_num_run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)


directory =os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Documents\\')
assert len(os.listdir(directory)) > 0, 'Download files in the folder'

new_documents_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'New Documents')
print(new_documents_dir)
if os.path.isdir(new_documents_dir):
    shutil.rmtree(new_documents_dir)
os.mkdir(new_documents_dir)

for filename in os.listdir(directory):
    if filename.endswith('.docx'):
        directory_file = os.path.join(directory, filename)
        new_document = NewDocuments(filename)
        main_text, links_text, headers_text, images_text = new_document.adding_margins()
        my_doc = docx2txt.process(directory_file, directory).split('\n')
        new_doc = docx.Document()
        section = new_doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Mm(main_settings.left_margin)
        section.right_margin = Mm(main_settings.right_margin)
        section.top_margin = Mm(main_settings.top_margin)
        section.bottom_margin = Mm(main_settings.bottom_margin)
        new_doc.styles['Normal'].font.name = main_settings.font_name
        new_doc.styles['Normal'].font.size = Pt(main_settings.pt)
        for phrase in new_document.text:
            if phrase == 'TABLE':
            elif phrase == 'IMAGE':
                new_document.adding_picture(images_text[0])
                del images_text[0]
            elif phrase == 'HEADER':
                new_document.adding_heading(headers_text[0])
                del headers_text[0]
            elif phrase == 'LINK':
            elif phrase == 'TEXT':
                new_document.adding_paragraph(main_text[0])
                del main_text[0]
            new_document.add_page_number(section.footer.add_paragraph().add)

end_time = time.time()
print(end_time - start_time)
